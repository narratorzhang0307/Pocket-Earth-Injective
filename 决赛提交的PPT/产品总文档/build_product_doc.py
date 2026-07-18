#!/usr/bin/env python3
"""Build the Pocket Earth on Injective master product document.

Design archetype: compact_reference_guide with an editorial cover.
The generated file stays fully editable and uses real Word headings, lists,
tables, hyperlinks, fields and captions.
"""

from __future__ import annotations

import re
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path("/Users/zhangcheng/Desktop/Pocket-Earth-Injective")
SOURCE = ROOT / "决赛提交的PPT/产品总文档/Pocket Earth on Injective-总产品文档-2026-07-18.md"
OUTPUT = ROOT / "决赛提交的PPT/产品总文档/Pocket Earth on Injective-总产品文档-2026-07-18.docx"

PAGE_WIDTH_DXA = 12240
PAGE_HEIGHT_DXA = 15840
MARGIN_DXA = 1440
CONTENT_WIDTH_DXA = PAGE_WIDTH_DXA - 2 * MARGIN_DXA

COLORS = {
    "navy": "0B2545",
    "blue": "2E74B5",
    "green": "2D6A55",
    "bright_green": "22B573",
    "cream": "F7F4EA",
    "light_blue": "E8EEF5",
    "light_green": "E7F2EC",
    "light_gray": "F2F4F7",
    "mid_gray": "6B7280",
    "dark": "171A1F",
    "orange": "D96732",
    "gold": "A67C2D",
    "white": "FFFFFF",
}

VISUALS = {
    "4. 产品总结构：一颗地球，四个相互连接的层": {
        "path": ROOT / "Agent 构架/pocket-earth-agent-architecture-page-2.png",
        "caption": "图 1｜Pocket Earth 的 Agent Harness、空间层与公共见证层",
        "width": 6.35,
    },
    "7. Public Earth：公共地球": {
        "path": ROOT / "work/public-map-poster-raw.png",
        "caption": "图 2｜Public Earth：公共知识便签在手机地图端展开，链上只保存公开承诺",
        "width": 3.05,
    },
    "10. Frost 身份卡、NFT 思想与品牌系统": {
        "path": ROOT / "design/frost-identity-concepts/codex-a-frost-collection-2026-07-18.png",
        "caption": "图 3｜Frost 形象系统：统一轮廓提供识别度，个体细节提供情感归属",
        "width": 6.2,
    },
}


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120) -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_cell_width(cell, width_dxa: int) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.find(qn("w:tcW"))
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width_dxa))
    tc_w.set(qn("w:type"), "dxa")


def set_table_borders(table, color="B8C0CC", size=6) -> None:
    tbl_pr = table._tbl.tblPr
    borders = tbl_pr.find(qn("w:tblBorders"))
    if borders is None:
        borders = OxmlElement("w:tblBorders")
        tbl_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = borders.find(qn(f"w:{edge}"))
        if tag is None:
            tag = OxmlElement(f"w:{edge}")
            borders.append(tag)
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    header = OxmlElement("w:tblHeader")
    header.set(qn("w:val"), "true")
    tr_pr.append(header)


def prevent_row_split(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    node = OxmlElement("w:cantSplit")
    tr_pr.append(node)


def set_run_font(run, size: float | None = None, bold=None, color=None, name="Calibri") -> None:
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)


def set_run_mono(run, size: float = 9.2) -> None:
    run.font.name = "Menlo"
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    run.font.size = Pt(size)
    run.font.color.rgb = RGBColor.from_string(COLORS["navy"])


def add_hyperlink(paragraph, text: str, url: str, color="2E74B5", underline=True):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    r_fonts = OxmlElement("w:rFonts")
    r_fonts.set(qn("w:ascii"), "Calibri")
    r_fonts.set(qn("w:hAnsi"), "Calibri")
    r_fonts.set(qn("w:eastAsia"), "PingFang SC")
    r_pr.append(r_fonts)
    c = OxmlElement("w:color")
    c.set(qn("w:val"), color)
    r_pr.append(c)
    if underline:
        u = OxmlElement("w:u")
        u.set(qn("w:val"), "single")
        r_pr.append(u)
    new_run.append(r_pr)
    txt = OxmlElement("w:t")
    txt.text = text
    new_run.append(txt)
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


INLINE_RE = re.compile(
    r"(\*\*.*?\*\*|`[^`]+`|\[[^\]]+\]\(https?://[^)]+\)|https?://[^\s，。；、)]+)"
)


def add_inline(paragraph, text: str, *, default_size=11.0, default_color=None) -> None:
    cursor = 0
    for match in INLINE_RE.finditer(text):
        if match.start() > cursor:
            run = paragraph.add_run(text[cursor:match.start()])
            set_run_font(run, default_size, color=default_color)
        token = match.group(0)
        if token.startswith("**") and token.endswith("**"):
            run = paragraph.add_run(token[2:-2])
            set_run_font(run, default_size, bold=True, color=default_color or COLORS["dark"])
        elif token.startswith("`") and token.endswith("`"):
            run = paragraph.add_run(token[1:-1])
            set_run_mono(run, default_size - 1)
            run.font.highlight_color = None
        elif token.startswith("["):
            m = re.match(r"\[([^\]]+)\]\((https?://[^)]+)\)", token)
            if m:
                add_hyperlink(paragraph, m.group(1), m.group(2))
        else:
            add_hyperlink(paragraph, token, token)
        cursor = match.end()
    if cursor < len(text):
        run = paragraph.add_run(text[cursor:])
        set_run_font(run, default_size, color=default_color)


def add_field(paragraph, code: str) -> None:
    run = paragraph.add_run()
    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = code
    separate = OxmlElement("w:fldChar")
    separate.set(qn("w:fldCharType"), "separate")
    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    run._r.extend([begin, instr, separate, end])


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    prefix = paragraph.add_run("POCKET EARTH · ")
    set_run_font(prefix, 8.5, color=COLORS["mid_gray"])
    add_field(paragraph, "PAGE")


def set_page_border(section, color="D7DCE3", size=8, offset=20) -> None:
    sect_pr = section._sectPr
    pg_borders = sect_pr.find(qn("w:pgBorders"))
    if pg_borders is None:
        pg_borders = OxmlElement("w:pgBorders")
        sect_pr.append(pg_borders)
    pg_borders.set(qn("w:offsetFrom"), "page")
    for edge in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{edge}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), str(size))
        tag.set(qn("w:space"), str(offset))
        tag.set(qn("w:color"), color)
        pg_borders.append(tag)


def configure_styles(doc: Document) -> None:
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(COLORS["dark"])
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    for name, size, color, before, after in (
        ("Title", 30, COLORS["navy"], 0, 14),
        ("Subtitle", 16, COLORS["green"], 0, 12),
        ("Heading 1", 17, COLORS["navy"], 18, 10),
        ("Heading 2", 13.5, COLORS["blue"], 14, 7),
        ("Heading 3", 11.5, COLORS["green"], 10, 5),
    ):
        style = styles[name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = RGBColor.from_string(color)
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True
        style.paragraph_format.line_spacing = 1.05

    if "Quote Callout" not in styles:
        quote = styles.add_style("Quote Callout", WD_STYLE_TYPE.PARAGRAPH)
    else:
        quote = styles["Quote Callout"]
    quote.font.name = "Calibri"
    quote._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    quote.font.size = Pt(11.5)
    quote.font.italic = False
    quote.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    quote.paragraph_format.left_indent = Inches(0.22)
    quote.paragraph_format.right_indent = Inches(0.12)
    quote.paragraph_format.space_before = Pt(5)
    quote.paragraph_format.space_after = Pt(9)
    quote.paragraph_format.line_spacing = 1.25

    if "Figure Caption" not in styles:
        caption = styles.add_style("Figure Caption", WD_STYLE_TYPE.PARAGRAPH)
    else:
        caption = styles["Figure Caption"]
    caption.font.name = "Calibri"
    caption._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    caption.font.size = Pt(9)
    caption.font.color.rgb = RGBColor.from_string(COLORS["mid_gray"])
    caption.paragraph_format.space_before = Pt(4)
    caption.paragraph_format.space_after = Pt(10)
    caption.paragraph_format.keep_with_next = False

    for style_name in ("List Bullet", "List Number"):
        style = styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
        style.font.size = Pt(11)
        style.paragraph_format.left_indent = Inches(0.375)
        style.paragraph_format.first_line_indent = Inches(-0.188)
        style.paragraph_format.space_after = Pt(4)
        style.paragraph_format.line_spacing = 1.25

    if "Manual Number" not in styles:
        manual = styles.add_style("Manual Number", WD_STYLE_TYPE.PARAGRAPH)
    else:
        manual = styles["Manual Number"]
    manual.font.name = "Calibri"
    manual._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    manual.font.size = Pt(11)
    manual.paragraph_format.left_indent = Inches(0.375)
    manual.paragraph_format.first_line_indent = Inches(-0.188)
    manual.paragraph_format.space_after = Pt(4)
    manual.paragraph_format.line_spacing = 1.25

    if "Code Block" not in styles:
        code = styles.add_style("Code Block", WD_STYLE_TYPE.PARAGRAPH)
    else:
        code = styles["Code Block"]
    code.font.name = "Menlo"
    code._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    code.font.size = Pt(8.6)
    code.font.color.rgb = RGBColor.from_string(COLORS["navy"])
    code.paragraph_format.left_indent = Inches(0.18)
    code.paragraph_format.right_indent = Inches(0.18)
    code.paragraph_format.space_before = Pt(6)
    code.paragraph_format.space_after = Pt(8)
    code.paragraph_format.line_spacing_rule = WD_LINE_SPACING.SINGLE


def shade_paragraph(paragraph, fill: str, border_color: str | None = None) -> None:
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)
    if border_color:
        p_bdr = p_pr.find(qn("w:pBdr"))
        if p_bdr is None:
            p_bdr = OxmlElement("w:pBdr")
            p_pr.append(p_bdr)
        left = OxmlElement("w:left")
        left.set(qn("w:val"), "single")
        left.set(qn("w:sz"), "20")
        left.set(qn("w:space"), "8")
        left.set(qn("w:color"), border_color)
        p_bdr.append(left)


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(24)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("POCKET EARTH")
    set_run_font(r, 11, bold=True, color=COLORS["bright_green"])
    r.font.letter_spacing = Pt(1.2) if hasattr(r.font, "letter_spacing") else None

    band = doc.add_table(rows=1, cols=1)
    band.alignment = WD_TABLE_ALIGNMENT.CENTER
    band.autofit = False
    set_table_borders(band, color=COLORS["navy"], size=0)
    cell = band.cell(0, 0)
    set_cell_width(cell, CONTENT_WIDTH_DXA)
    set_cell_margins(cell, top=300, bottom=300, start=220, end=220)
    set_cell_shading(cell, COLORS["navy"])
    title_p = cell.paragraphs[0]
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title = title_p.add_run("Pocket Earth on Injective")
    set_run_font(title, 30, bold=True, color=COLORS["white"])
    title_p.paragraph_format.space_after = Pt(8)
    sub_p = cell.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub = sub_p.add_run("总产品文档与决赛备忘录")
    set_run_font(sub, 16, bold=True, color="BFEBD7")

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(22)
    p.paragraph_format.space_after = Pt(18)
    add_inline(p, "私人记忆钉在地球 · 公共知识凝成版次 · Agent 身份由 Injective 见证", default_size=12.5, default_color=COLORS["navy"])

    meta = doc.add_table(rows=6, cols=2)
    meta.alignment = WD_TABLE_ALIGNMENT.CENTER
    meta.autofit = False
    set_table_borders(meta, color="D9DEE6", size=5)
    meta_rows = [
        ("版本", "2026-07-18 / 决赛冲刺版"),
        ("中文名", "口袋地球"),
        ("产品形态", "移动端优先 PWA + 空间 Agent 平台 + Injective 公共见证层 + Frost Edge Node"),
        ("代码仓库", "github.com/narratorzhang0307/Pocket-Earth-Injective"),
        ("决赛域名", "pocketearth-injective.throughtheglass.art"),
        ("用途", "PPT、口播、答辩、开发、验证与交接的共同母文档"),
    ]
    for row, (label, value) in zip(meta.rows, meta_rows):
        prevent_row_split(row)
        set_cell_width(row.cells[0], 1800)
        set_cell_width(row.cells[1], CONTENT_WIDTH_DXA - 1800)
        set_cell_shading(row.cells[0], COLORS["light_blue"])
        for cell in row.cells:
            set_cell_margins(cell, top=95, bottom=95, start=120, end=120)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        p0 = row.cells[0].paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        r0 = p0.add_run(label)
        set_run_font(r0, 9.5, bold=True, color=COLORS["navy"])
        p1 = row.cells[1].paragraphs[0]
        add_inline(p1, value, default_size=9.5, default_color=COLORS["dark"])

    quote = doc.add_paragraph(style="Quote Callout")
    quote.alignment = WD_ALIGN_PARAGRAPH.CENTER
    quote.paragraph_format.space_before = Pt(22)
    quote.paragraph_format.space_after = Pt(5)
    add_inline(quote, "“空间留在 Pocket Earth，时间由 Injective 见证。”", default_size=15, default_color=COLORS["green"])
    shade_paragraph(quote, COLORS["light_green"], COLORS["bright_green"])

    footer_note = doc.add_paragraph()
    footer_note.alignment = WD_ALIGN_PARAGRAPH.CENTER
    footer_note.paragraph_format.space_before = Pt(14)
    r = footer_note.add_run("MASTER PRODUCT DOCUMENT · INTERNAL & FINAL PITCH USE")
    set_run_font(r, 8.5, bold=True, color=COLORS["mid_gray"])
    doc.add_page_break()


def add_toc(doc: Document, source: str) -> None:
    h = doc.add_heading("目录", level=1)
    h.paragraph_format.page_break_before = False
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    add_inline(p, "以下为固定文档导航；打开 Word 后可右键自动目录 → 更新域 → 更新整个目录，以显示最新页码。", default_size=9.5, default_color=COLORS["mid_gray"])

    chapter_titles = [
        line[2:].strip() for line in source.splitlines()
        if re.match(r"^# \d+\. ", line)
    ]
    rows = (len(chapter_titles) + 1) // 2
    nav = doc.add_table(rows=rows, cols=2)
    nav.alignment = WD_TABLE_ALIGNMENT.CENTER
    nav.autofit = False
    set_table_borders(nav, color="D9DEE6", size=5)
    for idx in range(rows * 2):
        row_idx = idx % rows
        col_idx = idx // rows
        cell = nav.cell(row_idx, col_idx)
        set_cell_width(cell, CONTENT_WIDTH_DXA // 2)
        set_cell_margins(cell, top=70, bottom=70, start=110, end=110)
        if row_idx % 2 == 1:
            set_cell_shading(cell, COLORS["light_gray"])
        if idx < len(chapter_titles):
            pnav = cell.paragraphs[0]
            pnav.paragraph_format.space_after = Pt(0)
            add_inline(pnav, chapter_titles[idx], default_size=9.2, default_color=COLORS["navy"])

    auto = doc.add_paragraph()
    auto.paragraph_format.space_before = Pt(8)
    label = auto.add_run("自动目录（在 Word 中更新后显示页码）")
    set_run_font(label, 8.5, bold=True, color=COLORS["mid_gray"])
    toc = doc.add_paragraph()
    add_field(toc, 'TOC \\o "1-3" \\h \\z \\u')
    doc.add_page_break()


def split_table_row(line: str) -> list[str]:
    return [part.strip() for part in line.strip().strip("|").split("|")]


def table_widths(cols: int) -> list[int]:
    options = {
        2: [2500, CONTENT_WIDTH_DXA - 2500],
        3: [1900, 4700, CONTENT_WIDTH_DXA - 6600],
        4: [1650, 2950, 2550, CONTENT_WIDTH_DXA - 7150],
        5: [1450, 2050, 1900, 1800, CONTENT_WIDTH_DXA - 7200],
    }
    return options.get(cols, [CONTENT_WIDTH_DXA // cols] * cols)


def add_markdown_table(doc: Document, rows: list[list[str]]) -> None:
    cols = len(rows[0])
    widths = table_widths(cols)
    table = doc.add_table(rows=len(rows), cols=cols)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table)
    for r_idx, row_values in enumerate(rows):
        row = table.rows[r_idx]
        prevent_row_split(row)
        if r_idx == 0:
            set_repeat_table_header(row)
        for c_idx, value in enumerate(row_values):
            cell = row.cells[c_idx]
            set_cell_width(cell, widths[c_idx])
            set_cell_margins(cell)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            if r_idx == 0:
                set_cell_shading(cell, COLORS["navy"])
            elif r_idx % 2 == 0:
                set_cell_shading(cell, COLORS["light_gray"])
            p = cell.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.line_spacing = 1.12
            add_inline(
                p,
                value,
                default_size=8.5 if cols >= 4 else 9.2,
                default_color=COLORS["white"] if r_idx == 0 else COLORS["dark"],
            )
            if r_idx == 0:
                for run in p.runs:
                    run.bold = True
    after = doc.add_paragraph()
    after.paragraph_format.space_after = Pt(1)


def add_visual(doc: Document, heading_text: str) -> None:
    visual = VISUALS.get(heading_text)
    if not visual or not visual["path"].exists():
        return
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(0)
    p.add_run().add_picture(str(visual["path"]), width=Inches(visual["width"]))
    cap = doc.add_paragraph(style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run(visual["caption"])


def add_hardware_visuals(doc: Document) -> None:
    paths = [
        ROOT / "work/hardware-qa/root-final.png",
        ROOT / "work/hardware-qa/podcast-preview-final.png",
        ROOT / "work/hardware-qa/earth-answer-large.png",
    ]
    if not all(path.exists() for path in paths):
        return
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    set_table_borders(table, color="FFFFFF", size=0)
    for cell, path in zip(table.rows[0].cells, paths):
        set_cell_width(cell, CONTENT_WIDTH_DXA // 3)
        set_cell_margins(cell, top=30, bottom=30, start=45, end=45)
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(path), width=Inches(1.75))
    cap = doc.add_paragraph(style="Figure Caption")
    cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap.add_run("图 4｜Frost Edge Node：最外层启动器、口袋播客与地球答案")


def extract_body_lines(text: str) -> list[str]:
    lines = text.splitlines()
    start = next(i for i, line in enumerate(lines) if line.startswith("# 1. "))
    return lines[start:]


def add_body(doc: Document, lines: Iterable[str]) -> None:
    lines = list(lines)
    i = 0
    in_code = False
    code_lines: list[str] = []
    visual_inserted: set[str] = set()
    hardware_inserted = False

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.startswith("```"):
            if not in_code:
                in_code = True
                code_lines = []
            else:
                p = doc.add_paragraph(style="Code Block")
                shade_paragraph(p, COLORS["light_gray"], "CAD2DC")
                run = p.add_run("\n".join(code_lines))
                set_run_mono(run, 8.4)
                in_code = False
            i += 1
            continue
        if in_code:
            code_lines.append(line)
            i += 1
            continue

        if not line.strip():
            i += 1
            continue

        if line.startswith("# "):
            text = line[2:].strip()
            h = doc.add_heading(text, level=1)
            h.paragraph_format.page_break_before = True
            if text in VISUALS and text not in visual_inserted:
                add_visual(doc, text)
                visual_inserted.add(text)
            if text == "11. Frost Edge Node 实体终端" and not hardware_inserted:
                add_hardware_visuals(doc)
                hardware_inserted = True
            i += 1
            continue
        if line.startswith("## "):
            doc.add_heading(line[3:].strip(), level=2)
            i += 1
            continue
        if line.startswith("### "):
            doc.add_heading(line[4:].strip(), level=3)
            i += 1
            continue

        if line.startswith("|") and i + 1 < len(lines) and re.match(r"^\|?\s*:?-+", lines[i + 1].strip().lstrip("|")):
            table_lines = [line]
            i += 2  # skip separator line
            while i < len(lines) and lines[i].strip().startswith("|"):
                table_lines.append(lines[i].rstrip())
                i += 1
            add_markdown_table(doc, [split_table_row(x) for x in table_lines])
            continue

        if line.startswith(">"):
            quote_lines = []
            while i < len(lines) and lines[i].startswith(">"):
                quote_lines.append(lines[i][1:].strip().rstrip("  "))
                i += 1
            p = doc.add_paragraph(style="Quote Callout")
            add_inline(p, " ".join(quote_lines), default_size=11.5, default_color=COLORS["navy"])
            shade_paragraph(p, COLORS["light_green"], COLORS["bright_green"])
            continue

        bullet = re.match(r"^\s*[-*]\s+(.*)$", line)
        numbered = re.match(r"^\s*(\d+)\.\s+(.*)$", line)
        if bullet or numbered:
            if bullet:
                p = doc.add_paragraph(style="List Bullet")
                add_inline(p, bullet.group(1), default_size=11)
            else:
                p = doc.add_paragraph(style="Manual Number")
                number = p.add_run(f"{numbered.group(1)}. ")
                set_run_font(number, 11, bold=True, color=COLORS["navy"])
                add_inline(p, numbered.group(2), default_size=11)
            i += 1
            continue

        if line == "---":
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(5)
            p.paragraph_format.space_after = Pt(5)
            p_pr = p._p.get_or_add_pPr()
            p_bdr = OxmlElement("w:pBdr")
            bottom = OxmlElement("w:bottom")
            bottom.set(qn("w:val"), "single")
            bottom.set(qn("w:sz"), "6")
            bottom.set(qn("w:color"), "D9DEE6")
            p_bdr.append(bottom)
            p_pr.append(p_bdr)
            i += 1
            continue

        # Join wrapped Markdown lines into one paragraph, preserving hard breaks.
        para_lines = [line.rstrip("  ")]
        hard_breaks = [line.endswith("  ")]
        i += 1
        while i < len(lines):
            nxt = lines[i].rstrip()
            if not nxt.strip() or nxt.startswith(("#", ">", "|", "```")) or re.match(r"^\s*[-*]\s+", nxt) or re.match(r"^\s*\d+\.\s+", nxt):
                break
            para_lines.append(nxt.rstrip("  "))
            hard_breaks.append(nxt.endswith("  "))
            i += 1
        p = doc.add_paragraph()
        for idx, chunk in enumerate(para_lines):
            if idx:
                p.add_run().add_break(WD_BREAK.LINE if hard_breaks[idx - 1] else WD_BREAK.LINE)
            add_inline(p, chunk, default_size=11)


def configure_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(PAGE_WIDTH_DXA / 1440)
    section.page_height = Inches(PAGE_HEIGHT_DXA / 1440)
    section.orientation = WD_ORIENT.PORTRAIT
    section.top_margin = Inches(MARGIN_DXA / 1440)
    section.bottom_margin = Inches(MARGIN_DXA / 1440)
    section.left_margin = Inches(MARGIN_DXA / 1440)
    section.right_margin = Inches(MARGIN_DXA / 1440)
    section.header_distance = Inches(0.42)
    section.footer_distance = Inches(0.42)
    section.different_first_page_header_footer = True
    set_page_border(section)
    configure_styles(doc)

    settings = doc.settings._element
    update_fields = settings.find(qn("w:updateFields"))
    if update_fields is None:
        update_fields = OxmlElement("w:updateFields")
        settings.append(update_fields)
    update_fields.set(qn("w:val"), "true")

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    r = hp.add_run("POCKET EARTH ON INJECTIVE  ·  PRODUCT MASTER DOCUMENT")
    set_run_font(r, 8.2, bold=True, color=COLORS["mid_gray"])
    footer = section.footer
    add_page_number(footer.paragraphs[0])
    return doc


def add_document_properties(doc: Document) -> None:
    props = doc.core_properties
    props.title = "Pocket Earth on Injective — 总产品文档与决赛备忘录"
    props.subject = "产品、架构、链上证据、知识系统、硬件、Azure、决赛叙事"
    props.author = "Pocket Earth / Codex"
    props.keywords = "Pocket Earth, Injective, Agent, Public Earth, Merkle, Frost, Azure"
    props.comments = "2026-07-18 决赛冲刺主文档"


def build() -> Path:
    doc = configure_document()
    add_document_properties(doc)
    add_cover(doc)
    source = SOURCE.read_text(encoding="utf-8")
    add_toc(doc, source)
    add_body(doc, extract_body_lines(source))
    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    return OUTPUT


if __name__ == "__main__":
    print(build())
