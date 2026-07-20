from __future__ import annotations

import json
from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION_START
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


PROJECT = Path("/Users/zhangcheng/Desktop/Pocket-Earth-Injective")
PPT_DIR = PROJECT / "决赛提交的PPT"
SOURCE = PPT_DIR / "Pocket Earth on Injective-决赛路演PPT-完整40页-硬件Azure强化终稿-v2-第40页修正-2026-07-20.pptx.inspect.ndjson"
OUTPUT = PPT_DIR / "产品总文档" / "Pocket Earth on Injective-决赛路演PPT-40页逐页口播稿-2026-07-20.docx"


TITLES = [
    "Pocket Earth on Injective",
    "AI 时代，同时丢失了两件东西",
    "把信息炼成可验证知识",
    "把信息炼成一张知识地图",
    "Frost：一台努力理解人类的机器",
    "一张生活碎片，怎样长成属于你的地球",
    "三个入口，共用同一颗地球",
    "一颗地球，连接四个职责清楚的层",
    "让 Frost 逐步理解一个人",
    "当下保持连贯，长期只沉淀必要偏好",
    "总 Frost 负责判断，专业 Agent 负责完成",
    "每个人都可以创造适合 Pocket Earth 的 Agent",
    "公共地球在手机，公共事实在 Injective",
    "同一轮廓，五个公开分身",
    "一个 Harness，八个领域 Agent",
    "过程保留七天，精选版次长期存在",
    "正文留在资源包，版次根写入 Injective",
    "概率性的 Agent，需要确定性的公开锚点",
    "身份、门牌、握手与知识版次都能从产品追到交易",
    "一个模型入口，按任务选择合适能力",
    "手机负责发起与回看，Frost Edge 负责本地推理、现场反馈与安全执行",
    "树莓派内部的职责链",
    "从已运行真机原型，走向 Frost Edge 实体智能终端",
    "输入、确认与反馈形成完整物理闭环",
    "Frost Edge 硬件全貌",
    "一个 PI HOME，三种可现场运行的实体体验",
    "Gemma × Microsoft Foundry：端云双脑与可观察 Harness",
    "从代码到真机的可复现证据链",
    "让链上身份进入房间，公开见闻变成灯光、画面和声音",
    "只读公开证据，白名单动作驱动物理反馈",
    "音乐、知识与行动提示，共用一套清楚的桌面交互",
    "一条公开知识，走完整条证据与物理反馈路径",
    "把人生放回空间，把公共知识交给时间证明",
    "Azure 接入以真实请求为准，留下可复核的模型路由证据",
    "Frost 已经有一具能工作的身体",
    "一个 Launcher，收纳音乐、知识与行动提示",
    "服务器每天编排，树莓派只读同步同一版次",
    "事件合同只开放三类动作，其他数据到不了设备",
    "Discover → Verify → Approve → Commit → Read → Echo",
    "把人生放回空间，把公共知识交给时间证明",
]

CHAPTERS = [
    (1, 5, "开场与产品命题"),
    (6, 10, "私人地球与长期记忆"),
    (11, 20, "Agent 平台、公共知识与 Injective"),
    (21, 28, "Frost Edge 与 Microsoft Azure 端云协同"),
    (29, 38, "硬件交互、口袋播客与安全边界"),
    (39, 40, "现场演示与收束"),
]

GREEN = "21936D"
GREEN_DARK = "2A604F"
INK = "171918"
MUTED = "666D69"
SOFT = "8B918D"
MINT = "E2F1EA"
CREAM = "F7F2E6"
RULE = "C8CDCA"


def rgb(hex_value: str) -> RGBColor:
    return RGBColor.from_string(hex_value)


def set_run_font(run, ascii_font="Calibri", east_asia_font="PingFang SC", size=None, color=None, bold=None, italic=None):
    run.font.name = ascii_font
    run._element.get_or_add_rPr().rFonts.set(qn("w:ascii"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:hAnsi"), ascii_font)
    run._element.get_or_add_rPr().rFonts.set(qn("w:eastAsia"), east_asia_font)
    if size is not None:
        run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = rgb(color)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic


def set_cell_shading(paragraph, fill: str):
    p_pr = paragraph._p.get_or_add_pPr()
    shd = p_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        p_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_paragraph_border(paragraph, side: str, color: str, size: int = 14, space: int = 4):
    p_pr = paragraph._p.get_or_add_pPr()
    p_bdr = p_pr.find(qn("w:pBdr"))
    if p_bdr is None:
        p_bdr = OxmlElement("w:pBdr")
        p_pr.append(p_bdr)
    element = OxmlElement(f"w:{side}")
    element.set(qn("w:val"), "single")
    element.set(qn("w:sz"), str(size))
    element.set(qn("w:space"), str(space))
    element.set(qn("w:color"), color)
    p_bdr.append(element)


def set_keep(paragraph, keep_next=False, keep_lines=True):
    p_pr = paragraph._p.get_or_add_pPr()
    if keep_next:
        p_pr.append(OxmlElement("w:keepNext"))
    if keep_lines:
        p_pr.append(OxmlElement("w:keepLines"))


def add_page_number(paragraph):
    paragraph.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    run = paragraph.add_run()
    set_run_font(run, size=8.5, color=SOFT, bold=True)
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    result = OxmlElement("w:t")
    result.text = "1"
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")
    run._r.append(fld_begin)
    run._r.append(instr)
    run._r.append(fld_separate)
    run._r.append(result)
    run._r.append(fld_end)


def load_notes() -> list[str]:
    notes = {}
    for line in SOURCE.read_text(encoding="utf-8").splitlines():
        obj = json.loads(line)
        if obj.get("kind") == "notes":
            notes[int(obj["slide"])] = obj.get("text", "").strip()
    if sorted(notes) != list(range(1, 41)):
        raise RuntimeError(f"Expected slide notes 1..40, got {sorted(notes)}")
    return [notes[index] for index in range(1, 41)]


def configure_styles(doc: Document):
    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    normal.font.size = Pt(11)
    normal.font.color.rgb = rgb(INK)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(6)
    normal.paragraph_format.line_spacing = 1.25

    h1 = styles["Heading 1"]
    h1.font.name = "Calibri"
    h1._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h1._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    h1.font.size = Pt(16)
    h1.font.bold = True
    h1.font.color.rgb = rgb(GREEN_DARK)
    h1.paragraph_format.space_before = Pt(18)
    h1.paragraph_format.space_after = Pt(10)
    h1.paragraph_format.line_spacing = 1.0

    h2 = styles["Heading 2"]
    h2.font.name = "Calibri"
    h2._element.rPr.rFonts.set(qn("w:ascii"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:hAnsi"), "Calibri")
    h2._element.rPr.rFonts.set(qn("w:eastAsia"), "PingFang SC")
    h2.font.size = Pt(13)
    h2.font.bold = True
    h2.font.color.rgb = rgb(GREEN_DARK)
    h2.paragraph_format.space_before = Pt(14)
    h2.paragraph_format.space_after = Pt(7)
    h2.paragraph_format.line_spacing = 1.0


def configure_section(section):
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    section.top_margin = Inches(1)
    section.bottom_margin = Inches(1)
    section.left_margin = Inches(1)
    section.right_margin = Inches(1)
    section.header_distance = Inches(0.492)
    section.footer_distance = Inches(0.492)

    header = section.header
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hp.paragraph_format.space_after = Pt(3)
    left = hp.add_run("POCKET EARTH ON INJECTIVE")
    set_run_font(left, ascii_font="Arial", east_asia_font="PingFang SC", size=8.5, color=GREEN_DARK, bold=True)
    right = hp.add_run("    FINAL PITCH · 40 SLIDES")
    set_run_font(right, ascii_font="Arial", east_asia_font="PingFang SC", size=8.5, color=SOFT, bold=True)
    set_paragraph_border(hp, "bottom", RULE, size=6, space=5)

    footer = section.footer
    fp = footer.paragraphs[0]
    add_page_number(fp)


def add_cover(doc: Document):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(112)
    p.paragraph_format.space_after = Pt(18)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("FINAL PITCH SCRIPT")
    set_run_font(r, ascii_font="Arial", size=11, color=GREEN, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    r = p.add_run("Pocket Earth on Injective")
    set_run_font(r, size=30, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(26)
    r = p.add_run("决赛路演 PPT · 40 页逐页口播稿")
    set_run_font(r, size=15, color=GREEN_DARK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(54)
    r = p.add_run("与最终 PPT 演讲者备注逐页对应")
    set_run_font(r, size=10.5, color=MUTED, italic=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run("2026.07.20")
    set_run_font(r, ascii_font="Arial", size=11, color=INK, bold=True)

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Pocket Earth · Injective · Microsoft Azure · Frost Edge")
    set_run_font(r, ascii_font="Arial", size=9.5, color=SOFT)


def add_front_matter(doc: Document):
    doc.add_page_break()
    h = doc.add_paragraph("使用说明", style="Heading 1")
    set_keep(h, keep_next=True)
    p = doc.add_paragraph()
    set_cell_shading(p, MINT)
    set_paragraph_border(p, "left", GREEN, size=18, space=7)
    p.paragraph_format.left_indent = Inches(0.14)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(14)
    r = p.add_run("本文档原样提取最终 40 页 PPT 中的演讲者备注，页码与幻灯片一一对应。可直接用于逐页排练、压缩口播时长和现场提词。")
    set_run_font(r, size=11, color=GREEN_DARK, bold=True)

    h = doc.add_paragraph("章节索引", style="Heading 1")
    set_keep(h, keep_next=True)
    for start, end, label in CHAPTERS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(8)
        p.paragraph_format.left_indent = Inches(0.05)
        n = p.add_run(f"{start:02d}—{end:02d}  ")
        set_run_font(n, ascii_font="Arial", size=10.5, color=GREEN, bold=True)
        t = p.add_run(label)
        set_run_font(t, size=11.5, color=INK, bold=True)


def add_slide_entry(doc: Document, number: int, title: str, note: str):
    h = doc.add_paragraph(style="Heading 2")
    set_cell_shading(h, MINT)
    set_paragraph_border(h, "left", GREEN, size=18, space=7)
    h.paragraph_format.left_indent = Inches(0.12)
    h.paragraph_format.right_indent = Inches(0.06)
    set_keep(h, keep_next=True)
    r = h.add_run(f"SLIDE {number:02d}  ")
    set_run_font(r, ascii_font="Arial", size=10.5, color=GREEN, bold=True)
    r = h.add_run(title)
    set_run_font(r, size=13, color=GREEN_DARK, bold=True)

    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Inches(0.12)
    p.paragraph_format.right_indent = Inches(0.08)
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.keep_together = True
    p.paragraph_format.widow_control = True
    r = p.add_run(note)
    set_run_font(r, size=11, color=INK)


def main():
    notes = load_notes()
    if len(TITLES) != 40:
        raise RuntimeError(f"Expected 40 titles, got {len(TITLES)}")

    doc = Document()
    configure_styles(doc)
    configure_section(doc.sections[0])
    doc.core_properties.title = "Pocket Earth on Injective - 决赛路演 40 页逐页口播稿"
    doc.core_properties.subject = "最终 PPT 演讲者备注独立文档"
    doc.core_properties.author = "Pocket Earth"
    doc.core_properties.keywords = "Pocket Earth, Injective, Microsoft Azure, Frost Edge, 决赛路演"

    add_cover(doc)
    add_front_matter(doc)

    for chapter_index, (start, end, label) in enumerate(CHAPTERS, start=1):
        doc.add_page_break()
        h = doc.add_paragraph(f"{chapter_index:02d}  {label}", style="Heading 1")
        set_keep(h, keep_next=True)
        set_paragraph_border(h, "bottom", RULE, size=8, space=6)
        for number in range(start, end + 1):
            add_slide_entry(doc, number, TITLES[number - 1], notes[number - 1])

    OUTPUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(OUTPUT)
    print(OUTPUT)


if __name__ == "__main__":
    main()
